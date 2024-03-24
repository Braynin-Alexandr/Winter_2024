# -*- coding: utf-8 -*-

from time import time
from openpyxl import load_workbook, Workbook
from collections import Counter
from pprint import pprint
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


def get_test_mark(pupil_answers):
    test_points_marks={6:2, 11:3, 15:4}
    for test_point, mark in test_points_marks.items():
        if sum(pupil_answers)<=test_point: return mark
    else: return 5


def count_marks(marks):
    test_marks=Counter([mark[1] for mark in marks.values()])
    quarter_3_marks=Counter([mark[2] for mark in marks.values()])
    return test_marks, quarter_3_marks


def get_quality(marks):
    good_marks, bad_marks = (marks.get(4, 0) + marks.get(5, 0)), (marks.get(2, 0) + marks.get(3,0))
    return round((good_marks / (good_marks+bad_marks))* 100)


def get_school_performance(count_marks):
    performance_marks = 0
    for mark in range(3, 6): performance_marks+=count_marks.get(mark, 0)
    count_all_marks=sum(count_marks.values())
    return round((performance_marks / count_all_marks)* 100)


def get_avarage_mark(count_marks):
    sum_marks=sum([mark*count for mark, count in count_marks.items()])
    all_marks=sum(count_marks.values())
    return round(sum_marks/all_marks)

def get_avarage_task_done(pupils_answers):
    task_done=0
    for pupils_ans in pupils_answers: task_done+=len(list(filter(lambda point: point>0, pupils_ans)))
    return round(task_done/len(pupils_answers))

def get_marks_better(marks):
    marks_better=0
    for compare_marks in marks.values():
        if compare_marks[1]>compare_marks[2]: marks_better+=1
    return marks_better, (round ((marks_better/len(marks))*100))

def get_marks_worse(marks):
    marks_worse = 0
    for compare_marks in marks.values():
        if compare_marks[1] < compare_marks[2]: marks_worse += 1
    return marks_worse, (round((marks_worse / len(marks)) * 100))

def get_verification_results(pupil_present, pupil_absent, test_avarage_mark, quarter_3_avarage_mark):
    check_1=round((pupil_absent/ (len(pupil_present)+ pupil_absent))*100)
    if check_1 > 25: return 'результат недостоверный, более 25% учеников не явилось на экзамен'

    check_2 = abs(test_avarage_mark-quarter_3_avarage_mark)
    if check_2>0.5: return 'результат недостоверный, разница между средней оценкой за тест и средней оценкой за четверть превышает 0.5 баллов'

    pupils_3_mark=0
    for point in pupil_present.values():
        if point[0]==7: pupils_3_mark+=1
    check_3 = round((pupils_3_mark / len(pupil_present)) * 100)
    if check_3>25: return f'результат недостоверный, высокий процент учеников на нижней границе оценки'

def get_most_difficult_tasks(tables):
    critical_amount_pupils = round(len(tables)*0.3)
    total_tasks = len(tables[0])
    count_wrong_answers={num_tasks: 0 for num_tasks in range(1, total_tasks+1)}

    for table in tables:
        for num_task, answer in enumerate(table):
            if answer==0: count_wrong_answers[num_task+1]+=1

    most_difficult_tasks=tuple(task for task, number_wrong_answers in count_wrong_answers.items() if number_wrong_answers >= critical_amount_pupils)
    return '-' if len(most_difficult_tasks)==0 else most_difficult_tasks


info_from_excel = {}
pupils_test_answers=[]
pupil_absent=0

wb=load_workbook('Форма.xlsx')
ws=wb.active

for row in range(2, ws.max_row+1):
    pupil_name=ws.cell(row, 1).value
    pupil_answers = []
    for column in range(2, ws.max_column):
        answer = ws.cell(row, column).value
        if answer == 'отсутствовал':
            pupil_absent+=1
            break
        else:
            if answer is not None: pupil_answers.append(answer)
    else:
        pupils_test_answers.append(pupil_answers)
        pupil_test_mark=get_test_mark(pupil_answers)
        pupil_quarter_3_mark=ws.cell(row, ws.max_column).value
        info_from_excel[pupil_name] = (sum(pupil_answers),pupil_test_mark, pupil_quarter_3_mark)
wb.save('Форма.xlsx')

count_test_marks, count_quarter_3_marks = count_marks(info_from_excel)

test_quality = get_quality(count_test_marks)
quarter_3_quality = get_quality(count_quarter_3_marks)

test_performance = get_school_performance(count_test_marks)
quarter_3_performance = get_school_performance(count_quarter_3_marks)

test_avarage_mark=get_avarage_mark(count_test_marks)
quarter_3_avarage_mark=get_avarage_mark(count_quarter_3_marks)

test_avarage_task_done=get_avarage_task_done(pupils_test_answers)

pupil_better_marks = get_marks_better(info_from_excel)
pupil_worse_marks = get_marks_worse(info_from_excel)

verification_results=get_verification_results(info_from_excel, pupil_absent, test_avarage_mark, quarter_3_avarage_mark)

most_difficult_tasks=get_most_difficult_tasks(pupils_test_answers)


doc = Document()
doc.add_heading('Анализ ВПР').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('')

sentences = [
    'Дата проведения ВПР: ' ,
    'ФИО учителя: ',
    f'Класс: ',
    'Название предмета: ',
    f'Количество учеников класса, которые писали ВПР: {len(info_from_excel)}.',
    f'Количество учеников класса, которые не писали ВПР: {pupil_absent}.',
    f'Средний балл по предмету за 3-ю четверть: {quarter_3_avarage_mark}.',
    f'Средний балл за ВПР: {test_avarage_mark}.',
    f'Среднее кол-во решенных заданий: {test_avarage_task_done}.',
    f'Кол-во и % детей, повысивших свою оценку: {pupil_better_marks[0]}, ({pupil_better_marks[1]}%).',
    f'Кол-во и % детей, понизивших свою оценку: {pupil_worse_marks[0]}, ({pupil_worse_marks[1]}%).',
    f'Проверка достоверности результатов: {verification_results}.',
    f'Задачи, которые вызвали наибольшие трудности: {most_difficult_tasks}.'
]

for sentence in sentences:
    doc.add_paragraph(sentence)
doc.add_paragraph('')


table_1 = doc.add_paragraph('Список участников').alignment = WD_ALIGN_PARAGRAPH.CENTER
table_1 = doc.add_table(rows=len(info_from_excel)+1, cols=3)

headers_table_1 = ['Ученик', 'Баллы ВПР', 'Итоговая оцнка']
for i, header in enumerate(headers_table_1): table_1.cell(0, i).text = header

for row, (student_id, data) in enumerate(info_from_excel.items(), 1):
    table_1.cell(row, 0).text = str(student_id)  # ученик
    table_1.cell(row, 1).text = str(data[0])     # Баллы ВПР
    table_1.cell(row, 2).text = str(data[1])
doc.add_paragraph('')


table_2 = doc.add_paragraph('Сравнительный анализ основных поазателей').alignment = WD_ALIGN_PARAGRAPH.CENTER
table_2 = doc.add_table(rows=6, cols=3)
headers_table_2 = ['Оценка', '3-я четверть', 'ВПР']
for i, header in enumerate(headers_table_2): table_2.cell(0, i).text = header

marks = [5, 4, 3]
for i, mark in enumerate(marks, 1):
    table_2.cell(i, 0).text = str(mark)+':'
    table_2.cell(i, 1).text = str(count_quarter_3_marks.get(mark, '-'))
    table_2.cell(i, 2).text = str(count_test_marks.get(mark, '-'))

table_2.cell(4, 0).text = '% качества:'
table_2.cell(4, 1).text = str(quarter_3_quality)
table_2.cell(4, 2).text = str(test_quality)

table_2.cell(5, 0).text = '% успеваемости:'
table_2.cell(5, 1).text = str(quarter_3_performance)
table_2.cell(5, 2).text = str(test_performance)
doc.add_paragraph('')


# Сохранение документа
doc.save('тест.docx')
